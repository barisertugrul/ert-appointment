from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
PRO_ROOT = ROOT.parent / "ert-appointment-pro"

DOCX_FILES: dict[str, dict[str, Path]] = {
    "lite": {
        "tr": ROOT / "ERT-Randevu-Dokumantasyon-TR.docx",
        "en": ROOT / "ERT-Appointment-Documentation-EN.docx",
    },
    "pro": {
        "tr": PRO_ROOT / "ERT-Randevu-Pro-Dokumantasyon-TR.docx",
        "en": PRO_ROOT / "ERT-Appointment-Pro-Documentation-EN.docx",
    },
}

REPLACEMENTS = {
    'Arama kutusuna "Appointment Booking by ERT Pro" yazın': 'Arama kutusuna "Appointment Booking by ERT" yazın',
    'Search for "Appointment Booking by ERT Pro"': 'Search for "Appointment Booking by ERT"',
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Fix DOCX consistency issues while preserving document theme/layout.",
    )
    parser.add_argument(
        "--edition",
        choices=["lite", "pro", "all"],
        default="all",
        help="Target document edition(s).",
    )
    parser.add_argument(
        "--lang",
        choices=["tr", "en", "all"],
        default="all",
        help="Target language document(s).",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Analyze and report changes without saving files.",
    )
    return parser.parse_args()


def target_file_paths(edition: str, lang: str) -> list[Path]:
    editions = ["lite", "pro"] if edition == "all" else [edition]
    languages = ["tr", "en"] if lang == "all" else [lang]

    paths: list[Path] = []
    for current_edition in editions:
        for current_lang in languages:
            path = DOCX_FILES[current_edition][current_lang]
            if path.exists():
                paths.append(path)
            else:
                print(f"Skipped (missing): {path}")
    return paths


def replace_text(doc: Document) -> int:
    replacements_applied = 0

    for para in doc.paragraphs:
        original = para.text
        updated = original
        for src, dst in REPLACEMENTS.items():
            if src in updated:
                updated = updated.replace(src, dst)
        if updated != original:
            para.text = updated
            replacements_applied += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original = para.text
                    updated = original
                    for src, dst in REPLACEMENTS.items():
                        if src in updated:
                            updated = updated.replace(src, dst)
                    if updated != original:
                        para.text = updated
                        replacements_applied += 1

    return replacements_applied


def remove_adjacent_duplicate_paragraphs(doc: Document) -> int:
    removed = 0
    previous_normalized = None

    for para in list(doc.paragraphs):
        current = para.text.strip()
        normalized = current.lstrip("- ").strip()

        if not current:
            previous_normalized = None
            continue

        if previous_normalized and normalized == previous_normalized:
            element = para._element
            element.getparent().remove(element)
            removed += 1
            continue

        previous_normalized = normalized

    return removed


def main() -> None:
    args = parse_args()
    paths = target_file_paths(args.edition, args.lang)

    if not paths:
        print("No matching DOCX files found.")
        return

    for path in paths:
        doc = Document(path)
        replacement_count = replace_text(doc)
        duplicate_count = remove_adjacent_duplicate_paragraphs(doc)

        if not args.dry_run:
            doc.save(path)

        mode = "Checked" if args.dry_run else "Updated"
        print(
            f"{mode} {path.name} "
            f"(text replacements: {replacement_count}, adjacent duplicates removed: {duplicate_count})"
        )


if __name__ == "__main__":
    main()
