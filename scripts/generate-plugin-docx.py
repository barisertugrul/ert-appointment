from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
PRO_ROOT = ROOT.parent / "ert-appointment-pro"


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def build_lite_tr(path: Path) -> None:
    doc = Document()
    doc.add_heading("Appointment Booking by ERT (Lite) — Türkçe Dokümantasyon", level=0)
    doc.add_paragraph("Sürüm: 1.0.1 | WordPress: 6.2+ | PHP: 8.1+")

    doc.add_heading("1) Genel Bakış", level=1)
    doc.add_paragraph("WordPress için hızlı, mobil uyumlu ve esnek randevu rezervasyon eklentisi.")
    add_bullets(doc, [
        "Mode-bazlı akış: general, department_no_provider, department_with_provider, provider_only",
        "Form bazlı Booking Button Text override desteği",
        "Lite kanal: email | Pro kanallar: sms, whatsapp",
        "REST API ve shortcode desteği",
    ])

    doc.add_heading("2) Kurulum", level=1)
    add_bullets(doc, [
        "Eklentiyi /wp-content/plugins/ert-appointment klasörüne yükleyin.",
        "WordPress panelinden eklentiyi etkinleştirin.",
        "Ayarlar, sağlayıcılar, çalışma saatleri ve formları yapılandırın.",
        "Sayfaya [erta_booking] shortcode’unu ekleyin.",
    ])

    doc.add_heading("3) Bildirim Kanalları", level=1)
    add_table(doc, ["Kanal", "Sürüm", "Not"], [
        ["Email", "Lite", "Varsayılan kanal"],
        ["SMS", "Pro", "Twilio / NetGSM"],
        ["WhatsApp", "Pro", "Meta Cloud API"],
    ])

    doc.add_heading("4) 1.0.1 Değişiklik Özeti", level=1)
    add_bullets(doc, [
        "Mode-bazlı booking akışı ve dinamik adım görünürlüğü",
        "Personelsiz modlarda slot birleştirme ve provider_id koruma",
        "Form bazlı gönderim butonu metni desteği",
        "Bildirim kanal modelinin email/sms/whatsapp olarak genişletilmesi",
        "WP.org release/assets checklist ve üretim dokümanları",
    ])

    doc.add_heading("5) Doküman Referansları", level=1)
    add_bullets(doc, [
        "ERT-Appointment-Documentation.html",
        "readme.txt",
        "docs/wporg-release-checklist.md",
        "docs/wporg-svn-publish.md",
        "docs/wporg-assets-master-checklist.md",
    ])

    doc.save(path)


def build_lite_en(path: Path) -> None:
    doc = Document()
    doc.add_heading("Appointment Booking by ERT (Lite) — English Documentation", level=0)
    doc.add_paragraph("Version: 1.0.1 | WordPress: 6.2+ | PHP: 8.1+")

    doc.add_heading("1) Overview", level=1)
    doc.add_paragraph("A modern, mobile-friendly and flexible appointment booking plugin for WordPress.")
    add_bullets(doc, [
        "Mode-based flow: general, department_no_provider, department_with_provider, provider_only",
        "Per-form Booking Button Text override support",
        "Lite channel: email | Pro channels: sms, whatsapp",
        "REST API and shortcode support",
    ])

    doc.add_heading("2) Installation", level=1)
    add_bullets(doc, [
        "Upload the plugin to /wp-content/plugins/ert-appointment.",
        "Activate from the WordPress plugins screen.",
        "Configure settings, providers, working hours and forms.",
        "Add the [erta_booking] shortcode to a page.",
    ])

    doc.add_heading("3) Notification Channels", level=1)
    add_table(doc, ["Channel", "Edition", "Notes"], [
        ["Email", "Lite", "Default channel"],
        ["SMS", "Pro", "Twilio / NetGSM"],
        ["WhatsApp", "Pro", "Meta Cloud API"],
    ])

    doc.add_heading("4) 1.0.1 Change Summary", level=1)
    add_bullets(doc, [
        "Mode-based booking flow with dynamic step visibility",
        "Provider-pool slot aggregation for personless modes",
        "Per-form submit button text override",
        "Notification model expanded to email/sms/whatsapp",
        "WP.org release/assets checklists and production guides",
    ])

    doc.add_heading("5) Documentation References", level=1)
    add_bullets(doc, [
        "ERT-Appointment-Documentation.html",
        "readme.txt",
        "docs/wporg-release-checklist.md",
        "docs/wporg-svn-publish.md",
        "docs/wporg-assets-master-checklist.md",
    ])

    doc.save(path)


def build_pro_tr(path: Path) -> None:
    doc = Document()
    doc.add_heading("Appointment Booking by ERT Pro — Türkçe Dokümantasyon", level=0)
    doc.add_paragraph("Sürüm: 1.0.0+ | Bağımlılık: Lite eklenti aktif olmalı")

    doc.add_heading("1) Pro Genel Bakış", level=1)
    add_bullets(doc, [
        "Ödeme altyapısı (PayTR, İyzico, Stripe, PayPal)",
        "Google Calendar senkronizasyonu",
        "Zoom toplantı otomasyonu",
        "SMS (Twilio/NetGSM) ve WhatsApp (Meta Cloud API)",
        "Raporlar ve waitlist yönetimi",
    ])

    doc.add_heading("2) Pro Bildirim Kanalları", level=1)
    add_table(doc, ["Kanal", "Sağlayıcı", "Kapsam"], [
        ["SMS", "Twilio / NetGSM", "Müşteri/sağlayıcı bildirimleri"],
        ["WhatsApp", "Meta Cloud API", "Müşteri odaklı şablonlar + hatırlatmalar"],
    ])

    doc.add_heading("3) Kurulum", level=1)
    add_bullets(doc, [
        "Lite eklentinin aktif olduğundan emin olun.",
        "Pro eklentiyi yükleyip etkinleştirin.",
        "Lisans anahtarını girip doğrulayın.",
        "Entegrasyonlar (Google/Zoom/SMS/WhatsApp) ayarlarını tamamlayın.",
    ])

    doc.add_heading("4) Referans Dokümanlar", level=1)
    add_bullets(doc, [
        "ERT-Appointment-Pro-Documentation.html",
        "README.md (pro)",
        "CHANGELOG.md (pro)",
    ])

    doc.save(path)


def build_pro_en(path: Path) -> None:
    doc = Document()
    doc.add_heading("Appointment Booking by ERT Pro — English Documentation", level=0)
    doc.add_paragraph("Version: 1.0.0+ | Dependency: Lite plugin must be active")

    doc.add_heading("1) Pro Overview", level=1)
    add_bullets(doc, [
        "Payment stack (PayTR, Iyzico, Stripe, PayPal)",
        "Google Calendar synchronization",
        "Zoom meeting automation",
        "SMS (Twilio/NetGSM) and WhatsApp (Meta Cloud API)",
        "Reports and waitlist management",
    ])

    doc.add_heading("2) Pro Notification Channels", level=1)
    add_table(doc, ["Channel", "Provider", "Scope"], [
        ["SMS", "Twilio / NetGSM", "Customer/provider notifications"],
        ["WhatsApp", "Meta Cloud API", "Customer templates and reminders"],
    ])

    doc.add_heading("3) Installation", level=1)
    add_bullets(doc, [
        "Ensure Lite plugin is active.",
        "Install and activate Pro plugin.",
        "Enter and validate license key.",
        "Configure Google/Zoom/SMS/WhatsApp integrations.",
    ])

    doc.add_heading("4) Reference Docs", level=1)
    add_bullets(doc, [
        "ERT-Appointment-Pro-Documentation.html",
        "README.md (pro)",
        "CHANGELOG.md (pro)",
    ])

    doc.save(path)


def main() -> None:
    build_lite_tr(ROOT / "ERT-Randevu-Dokumantasyon-TR.docx")
    build_lite_en(ROOT / "ERT-Appointment-Documentation-EN.docx")

    PRO_ROOT.mkdir(parents=True, exist_ok=True)
    build_pro_tr(PRO_ROOT / "ERT-Randevu-Pro-Dokumantasyon-TR.docx")
    build_pro_en(PRO_ROOT / "ERT-Appointment-Pro-Documentation-EN.docx")

    print("DOCX documentation files generated successfully.")


if __name__ == "__main__":
    main()
