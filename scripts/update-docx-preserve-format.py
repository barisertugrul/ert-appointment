from pathlib import Path
from shutil import copy2
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
PRO_ROOT = ROOT.parent / "ert-appointment-pro"


def replace_in_runs(paragraph, replacements: dict[str, str]) -> None:
    for run in paragraph.runs:
        text = run.text
        for src, dst in replacements.items():
            if src in text:
                text = text.replace(src, dst)
        run.text = text


def apply_replacements(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        replace_in_runs(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_runs(paragraph, replacements)


def add_heading_safe(doc: Document, text: str) -> None:
    try:
        doc.add_heading(text, level=1)
    except KeyError:
        paragraph = doc.add_paragraph(text)
        if paragraph.runs:
            paragraph.runs[0].bold = True


def add_bullet_safe(doc: Document, text: str) -> None:
    try:
        doc.add_paragraph(text, style="List Bullet")
    except KeyError:
        doc.add_paragraph(f"- {text}")


def add_appendix(doc: Document, title: str, bullets: list[str]) -> None:
    doc.add_page_break()
    add_heading_safe(doc, title)
    for bullet in bullets:
        add_bullet_safe(doc, bullet)


def update_lite_docs() -> None:
    tr_path = ROOT / "ERT-Randevu-Dokumantasyon-TR.docx"
    en_path = ROOT / "ERT-Appointment-Documentation-EN.docx"

    tr = Document(tr_path)
    apply_replacements(tr, {
        "1.0.0": "1.0.1",
        "Version 1.0": "Version 1.0.1",
        "Sürüm 1.0": "Sürüm 1.0.1",
        "WordPress 6.0+": "WordPress 6.2+",
        "WordPress 6.0": "WordPress 6.2",
        "Slot Süresi": "Randevu Süresi",
        "E-posta Bildirimleri": "Bildirim Şablonları",
    })
    add_appendix(
        tr,
        "v1.0.1 Güncelleme Notları (Format Korunarak)",
        [
            "Mode-bazlı booking akışı: general, department_no_provider, department_with_provider, provider_only.",
            "Form bazlı Booking Button Text override desteği eklendi.",
            "Bildirim kanalları modeli genişletildi: Lite email, Pro sms/whatsapp.",
            "Pro WhatsApp provider (Meta Cloud API) ve SMS sağlayıcıları (Twilio/NetGSM) desteklenir.",
            "WP.org release, SVN publish ve assets üretim dokümanları eklendi.",
        ],
    )
    tr.save(tr_path)

    en = Document(en_path)
    apply_replacements(en, {
        "1.0.0": "1.0.1",
        "Version 1.0": "Version 1.0.1",
        "WordPress 6.0+": "WordPress 6.2+",
        "WordPress 6.0": "WordPress 6.2",
        "Slot Duration": "Appointment Duration",
        "Email Notifications": "Notification Templates",
    })
    add_appendix(
        en,
        "v1.0.1 Update Notes (Theme Preserved)",
        [
            "Mode-based booking flow: general, department_no_provider, department_with_provider, provider_only.",
            "Per-form Booking Button Text override support added.",
            "Notification channel model expanded: Lite email, Pro sms/whatsapp.",
            "Pro WhatsApp provider (Meta Cloud API) and SMS providers (Twilio/NetGSM) supported.",
            "WP.org release, SVN publish and assets documentation added.",
        ],
    )
    en.save(en_path)


def build_pro_docs_from_lite_theme() -> None:
    PRO_ROOT.mkdir(parents=True, exist_ok=True)

    src_tr = ROOT / "ERT-Randevu-Dokumantasyon-TR.docx"
    src_en = ROOT / "ERT-Appointment-Documentation-EN.docx"
    dst_tr = PRO_ROOT / "ERT-Randevu-Pro-Dokumantasyon-TR.docx"
    dst_en = PRO_ROOT / "ERT-Appointment-Pro-Documentation-EN.docx"

    copy2(src_tr, dst_tr)
    copy2(src_en, dst_en)

    tr = Document(dst_tr)
    apply_replacements(tr, {
        "Appointment Booking by ERT": "Appointment Booking by ERT Pro",
        "Lite (Ücretsiz)": "Lite (Ücretsiz) + Pro (Aktif)",
    })
    add_appendix(
        tr,
        "Pro Eklenti Katmanı — Ek Yetenekler",
        [
            "Ödeme: PayTR, İyzico, Stripe, PayPal.",
            "Entegrasyonlar: Google Calendar, Zoom.",
            "Bildirimler: SMS (Twilio/NetGSM), WhatsApp (Meta Cloud API).",
            "Raporlar ve waitlist modülleri.",
            "Not: Pro katmanı, Lite eklenti aktif olmadan çalışmaz.",
        ],
    )
    tr.save(dst_tr)

    en = Document(dst_en)
    apply_replacements(en, {
        "Appointment Booking by ERT": "Appointment Booking by ERT Pro",
        "Lite (Free)": "Lite (Free) + Pro (Active)",
    })
    add_appendix(
        en,
        "Pro Add-on Layer — Additional Capabilities",
        [
            "Payments: PayTR, Iyzico, Stripe, PayPal.",
            "Integrations: Google Calendar, Zoom.",
            "Notifications: SMS (Twilio/NetGSM), WhatsApp (Meta Cloud API).",
            "Reports and waitlist modules.",
            "Note: Pro layer requires the Lite plugin to be active.",
        ],
    )
    en.save(dst_en)


def main() -> None:
    update_lite_docs()
    build_pro_docs_from_lite_theme()
    print("DOCX files updated with preserved structure/theme.")


if __name__ == "__main__":
    main()
